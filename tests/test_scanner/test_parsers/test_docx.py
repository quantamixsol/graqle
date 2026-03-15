"""Tests for graqle.scanner.parsers.docx — DOCX parser."""

# ── graqle:intelligence ──
# module: tests.test_scanner.test_parsers.test_docx
# risk: LOW (impact radius: 0 modules)
# dependencies: __future__, pathlib, pytest, docx
# constraints: none
# ── /graqle:intelligence ──

from __future__ import annotations

from pathlib import Path

import pytest

docx = pytest.importorskip("docx", reason="python-docx not installed")

from graqle.scanner.parsers.docx import DOCXParser


@pytest.fixture
def parser() -> DOCXParser:
    return DOCXParser()


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX file."""
    from docx import Document
    from docx.shared import Inches

    p = tmp_path / "test.docx"
    doc = Document()
    doc.core_properties.title = "Test Document"
    doc.core_properties.author = "Test Author"

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction paragraph.")

    doc.add_heading("Details", level=2)
    doc.add_paragraph("Here are some details about the project.")
    doc.add_paragraph("The auth_service handles authentication.")

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "API"
    table.cell(1, 1).text = "REST"
    table.cell(2, 0).text = "Auth"
    table.cell(2, 1).text = "JWT"

    doc.save(str(p))
    return p


@pytest.fixture
def empty_docx(tmp_path: Path) -> Path:
    """Create an empty DOCX file."""
    from docx import Document

    p = tmp_path / "empty.docx"
    doc = Document()
    doc.save(str(p))
    return p


class TestDOCXParser:
    def test_is_available(self, parser: DOCXParser) -> None:
        assert parser.is_available() is True

    def test_supported_extensions(self, parser: DOCXParser) -> None:
        assert ".docx" in parser.supported_extensions

    def test_parse_basic(self, parser: DOCXParser, sample_docx: Path) -> None:
        doc = parser.parse(sample_docx)
        assert doc.format == "docx"
        assert doc.path == sample_docx
        assert len(doc.sections) >= 2
        assert doc.full_text  # non-empty

    def test_parse_title(self, parser: DOCXParser, sample_docx: Path) -> None:
        doc = parser.parse(sample_docx)
        assert doc.title == "Test Document"

    def test_parse_metadata(self, parser: DOCXParser, sample_docx: Path) -> None:
        doc = parser.parse(sample_docx)
        assert doc.metadata.get("author") == "Test Author"
        assert doc.metadata.get("title") == "Test Document"

    def test_parse_headings(self, parser: DOCXParser, sample_docx: Path) -> None:
        doc = parser.parse(sample_docx)
        heading_sections = [s for s in doc.sections if s.section_type == "heading"]
        assert len(heading_sections) >= 2

    def test_parse_tables(self, parser: DOCXParser, sample_docx: Path) -> None:
        doc = parser.parse(sample_docx)
        # Tables should be extracted
        all_tables = []
        for section in doc.sections:
            all_tables.extend(section.tables)
        assert len(all_tables) >= 1

    def test_parse_empty(self, parser: DOCXParser, empty_docx: Path) -> None:
        doc = parser.parse(empty_docx)
        assert doc.format == "docx"
        # May have 0 sections if empty

    def test_file_not_found(self, parser: DOCXParser, tmp_path: Path) -> None:
        with pytest.raises(FileNotFoundError):
            parser.parse(tmp_path / "nonexistent.docx")

    def test_content_in_full_text(self, parser: DOCXParser, sample_docx: Path) -> None:
        doc = parser.parse(sample_docx)
        assert "auth_service" in doc.full_text
