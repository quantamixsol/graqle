"""DOCX document parser — requires ``python-docx``.

Extracts text, headings, tables, and metadata from DOCX files.
"""

# ── graqle:intelligence ──
# module: graqle.scanner.parsers.docx
# risk: LOW (impact radius: 1 modules)
# consumers: test_docx
# dependencies: __future__, pathlib, typing, base
# constraints: none
# ── /graqle:intelligence ──

from __future__ import annotations

from pathlib import Path
from typing import Any

from graqle.scanner.parsers.base import (
    BaseDocParser,
    ParsedDocument,
    ParsedSection,
)


class DOCXParser(BaseDocParser):
    """Parse ``.docx`` files using ``python-docx``."""

    def parse(self, path: Path) -> ParsedDocument:
        """Parse a DOCX file into a :class:`ParsedDocument`."""
        import docx

        path = Path(path)
        if not path.is_file():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        metadata: dict[str, Any] = {"source": str(path)}
        parse_errors: list[str] = []

        try:
            doc = docx.Document(str(path))
        except Exception as exc:
            raise ValueError(f"Failed to parse DOCX: {exc}") from exc

        # Extract metadata
        core = doc.core_properties
        if core.title:
            metadata["title"] = core.title
        if core.author:
            metadata["author"] = core.author
        if core.created:
            metadata["created"] = str(core.created)
        if core.modified:
            metadata["modified"] = str(core.modified)

        # Extract tables
        all_tables: list[dict[str, Any]] = []
        for table in doc.tables:
            try:
                rows_data = []
                for row in table.rows:
                    rows_data.append([cell.text.strip() for cell in row.cells])
                if len(rows_data) >= 2:
                    headers = rows_data[0]
                    rows = rows_data[1:]
                    all_tables.append({"headers": headers, "rows": rows})
            except Exception as exc:
                parse_errors.append(f"Table extraction: {exc}")

        # Build sections from paragraphs
        sections: list[ParsedSection] = []
        current_title = ""
        current_level = 1
        current_content: list[str] = []
        current_tables: list[dict[str, Any]] = []

        _HEADING_STYLES = {
            "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
            "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
        }

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            level = _HEADING_STYLES.get(style_name)

            if level is not None:
                # Flush current section
                if current_content or current_title:
                    sections.append(
                        ParsedSection(
                            title=current_title,
                            content="\n".join(current_content).strip(),
                            level=current_level,
                            section_type="heading" if current_title else "paragraph",
                            tables=current_tables,
                            code_blocks=[],
                            links=[],
                        )
                    )
                current_title = para.text.strip()
                current_level = level
                current_content = []
                current_tables = []
            else:
                text = para.text.strip()
                if text:
                    current_content.append(text)

        # Flush final section
        if current_content or current_title:
            sections.append(
                ParsedSection(
                    title=current_title,
                    content="\n".join(current_content).strip(),
                    level=current_level,
                    section_type="heading" if current_title else "paragraph",
                    tables=current_tables,
                    code_blocks=[],
                    links=[],
                )
            )

        # Attach tables to the first section (or create a standalone section)
        if all_tables:
            if sections:
                sections[0] = ParsedSection(
                    title=sections[0].title,
                    content=sections[0].content,
                    level=sections[0].level,
                    section_type=sections[0].section_type,
                    tables=sections[0].tables + all_tables,
                    code_blocks=sections[0].code_blocks,
                    links=sections[0].links,
                )
            else:
                sections.append(
                    ParsedSection(
                        title="",
                        content="",
                        level=1,
                        section_type="paragraph",
                        tables=all_tables,
                        code_blocks=[],
                        links=[],
                    )
                )

        title = metadata.get("title", path.stem)
        full_text = "\n\n".join(
            (f"{s.title}\n{s.content}" if s.title else s.content)
            for s in sections
        ).strip()

        return ParsedDocument(
            path=path,
            title=str(title),
            format="docx",
            sections=sections,
            full_text=full_text,
            metadata=metadata,
            parse_errors=parse_errors,
        )

    def is_available(self) -> bool:
        try:
            import docx  # noqa: F401
            return True
        except ImportError:
            return False

    def missing_dependency_message(self) -> str:
        return "Install python-docx: pip install graqle[docs]"

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]
